import streamlit as st
import os
from base64 import b64encode
from dotenv import load_dotenv
from PIL import Image
import openai
import datetime
import locale
import json
import pandas as pd

PASTA_LAUDOS = "relatorios"

# Carregar orientações da estrutura do laudo
# Carrega arquivo .env
load_dotenv('.env')
client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'))


def carregar_orientacoes_estrutura():
    """Carrega as orientações para a estrutura do laudo"""
    try:
        if os.path.exists("orientacoes_estrutura_laudo.json"):
            with open("orientacoes_estrutura_laudo.json", "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as e:
        st.error(f"❌ Erro ao carregar orientações da estrutura: {str(e)}")
    return None


def listar_relatorios():
    """Retorna a lista de pastas de relatórios disponíveis."""
    return [nome for nome in os.listdir(PASTA_LAUDOS) if os.path.isdir(os.path.join(PASTA_LAUDOS, nome))]


def verificar_relatorio_existente(relatorio_nome):
    """Verifica se o relatório já foi gerado e retorna informações sobre ele."""
    pasta_relatorio = os.path.join(PASTA_LAUDOS, relatorio_nome)
    caminho_txt = os.path.join(pasta_relatorio, "relatorio_gerado.txt")
    caminho_docx = os.path.join(pasta_relatorio, "relatorio_gerado.docx")

    relatorio_existe = os.path.exists(
        caminho_txt) or os.path.exists(caminho_docx)

    if relatorio_existe:
        info_relatorio = {}
        if os.path.exists(caminho_txt):
            info_relatorio['txt'] = {
                'existe': True,
                'tamanho': os.path.getsize(caminho_txt),
                'data_modificacao': datetime.datetime.fromtimestamp(os.path.getmtime(caminho_txt))
            }
        if os.path.exists(caminho_docx):
            info_relatorio['docx'] = {
                'existe': True,
                'tamanho': os.path.getsize(caminho_docx),
                'data_modificacao': datetime.datetime.fromtimestamp(os.path.getmtime(caminho_docx))
            }
        return True, info_relatorio
    return False, {}


def carregar_dados_relatorio(relatorio_nome):
    """Carrega todos os dados necessários para o relatório"""
    try:
        # Verificar se o arquivo Excel existe
        if not os.path.exists("historico_relatorios.xlsx"):
            return None, "Arquivo historico_relatorios.xlsx não encontrado.", None

        # Carregar dados do Excel
        df = pd.read_excel("historico_relatorios.xlsx")
        dados = df[df["ID Relatório"] == relatorio_nome]

        if dados.empty:
            return None, f"Dados do relatório {relatorio_nome} não encontrados no Excel.", None

        # Converter para Series para evitar problemas de DataFrame
        dados = dados.iloc[0]

        # Carregar descrições e ditados
        pasta_relatorio = os.path.join(PASTA_LAUDOS, relatorio_nome)
        caminho_descricoes = os.path.join(pasta_relatorio, "descricoes.json")

        descricoes = {}
        if os.path.exists(caminho_descricoes):
            try:
                with open(caminho_descricoes, "r", encoding="utf-8") as f:
                    descricoes = json.load(f)
            except Exception as e:
                st.warning(f"Aviso: Erro ao carregar descrições: {str(e)}")

        return dados, descricoes, None
    except Exception as e:
        return None, None, f"Erro ao carregar dados: {str(e)}"


def gerar_preambulo(dados):
    """Gera o preâmbulo do relatório (texto inicial)"""
    try:
        locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
    except:
        try:
            locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil')
        except:
            pass

    try:
        data_pericia_dt = pd.to_datetime(dados["Data da Perícia"])
        dia = data_pericia_dt.day
        mes_extenso = data_pericia_dt.strftime('%B')
        ano = data_pericia_dt.year
        data_pericia_extenso = f"{dia} dias do mês de {mes_extenso} de {ano}"
    except Exception as e:
        data_pericia_extenso = str(dados["Data da Perícia"])

    preambulo = f"""
Aos {data_pericia_extenso}, nesta cidade de São Paulo e no Instituto de Criminalística "Perito Criminal Dr. Octávio Eduardo de Brito Alvarenga", da Superintendência da Polícia Técnico-Científica, da Secretaria da Segurança Pública, em conformidade com o disposto no artigo 178 do Decreto-Lei n. 3.689, de 03 de Outubro de 1941; pelo Senhor Diretor deste I.C., foi designado o Perito Criminal DR. {dados["Nome do Perito"]} para proceder ao exame supra-especificado, em atendimento à requisição do {dados["Requisitante"]} do {dados.get("Destinatário", "")}, objeto do B.O. n {dados["Número do BO"]}.
"""
    return preambulo


def gerar_historico(dados, orientacoes):
    """
    Gera a seção de histórico da requisição, reescrevendo o campo 'Histórico da Requisição'
    do Excel de forma mais clara e objetiva, utilizando a IA.
    """
    historico_requisicao = dados.get("Histórico da Requisição", "")

    if not historico_requisicao:
        return "## HISTÓRICO\n\nHistórico da requisição não informado."

    prompt = (
        f"Reescreva o texto abaixo, que é o histórico da requisição de perícia, de forma mais clara, objetiva e adequada para um laudo pericial. "
        f"Seja fiel às informações, mas melhore a redação e a estrutura."
        f"Histórico original:\n{historico_requisicao}"
    )

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "Você é um perito criminal experiente. Sua tarefa é reescrever o histórico da requisição "
                    "de perícia para que fique mais claro, objetivo e tecnicamente adequado ao laudo pericial."
                ),
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                ],
            },
        ],
        max_tokens=200,
    )
    historico_processado = response.choices[0].message.content.strip()

    historico_final = f"""## HISTÓRICO e DADOS VEICULARES
    {historico_processado}
    \n
    Referente ao veículo analisado trata-se da marca "{dados.get('Marca', '')}", modelo "{dados.get('Modelo', '')}", cor "{dados.get('Cor', '')}". Os pneus encontram-se em condição "{dados.get('Pneus', '')}" e os sistemas veiculares estão "{dados.get('Sistemas', '')}".
    """
    return historico_final


def gerar_objetivo(dados, orientacoes):
    """
    Retorna o texto padrão para a seção de objetivo do laudo pericial.
    O campo objetivo é sempre este texto fixo.
    """
    objetivo = "## Objetivo\n\nO presente exame pericial tem por objetivo atender a requisição de exame expedida pela Autoridade Policial, a fim de responder aos quesitos formulados."
    return objetivo


def gerar_quesitos(dados):
    """Gera a seção de quesitos copiando exatamente o campo Quesito da Perícia"""
    quesito_pericia = dados.get("Quesito da Perícia", "")

    if not quesito_pericia:
        return "## QUESITOS\n\nQuesitos da perícia não informados."

    quesitos = f"""
## QUESITOS

{quesito_pericia}
"""
    return quesitos


def gerar_exames(descricoes, orientacoes):
    """Gera a seção de exames com fotos, descrições e ditados periciais"""
    # Usar as orientações como prompt para estruturar a seção
    prompt_exames = orientacoes.get("exames", "") if orientacoes else ""

    exames = "## -------------------- EXAMES REALIZADOS ---------------------------- \n\n"

    # Processar imagens e descrições
    imagens_processadas = []
    for chave, valor in descricoes.items():
        if chave != "ditado_pericia":
            imagens_processadas.append((chave, valor))

    if imagens_processadas:
        exames += "### Documentação Fotográfica e Técnica\n\n"
        for imagem, descricao in imagens_processadas:
            exames += f"**Imagem:** {imagem}\n"
            exames += f"**Descrição:** {descricao}\n\n ******* FIM DO EXAME DESTA IMAGEM ******* \n\n"

    return exames


def gerar_conclusao(dados, orientacoes):
    """
    Gera a seção de conclusão respondendo aos quesitos.
    Utiliza informações do histórico, ditado pericial e descrições.
    """
    historico_requisicao = dados.get("Histórico da Requisição", "")
    ditado_pericial = dados.get("Ditado Pericial", "")
    descricoes = dados.get("Descrições", "")
    quesitos = dados.get("Quesito da Perícia", "")

    prompt = (
        "Vamos concluir o laudo pericial com as informações que já foram inseridas. "
        "Você deve se comportar como um perito e se basear apenas nos dados que foram apresentados anteriormente. "
        "Utilize como base as descrições das fotografias, o ditado pericial e o histórico da requisição para formular se o que foi dito no histórico da requisição é plausível ou não.\n"
        f"Histórico original:\n{historico_requisicao}\n"
        f"Ditado pericial:\n{ditado_pericial}\n"
        f"Descrições:\n{descricoes}\n"
        f"Quesitos:\n{quesitos}\n"
    )

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "Você é um perito criminal experiente. Sua tarefa é concluir o relatório pericial com base apenas nas informações fornecidas anteriormente. Seja claro, simples e objetivo.\n"
                    "Responda obrigatoriamente aos quesitos apresentados no contexto. Não inclua informações que não estejam nos dados fornecidos.\n"
                    "Na conclusão, inclua obrigatoriamente a orientação dos danos como uma composição de dois vetores, escolhendo e mencionando duas das opções abaixo (o dano é sempre bidimensional):\n"
                    "  (1) da esquerda para a direita OU da direita para a esquerda;\n"
                    "  (2) de cima para baixo OU de baixo para cima;\n"
                    "  (3) de frente para trás OU de trás para frente.\n"
                    "Considere sempre a perspectiva do motorista do veículo para definir os vetores.\n"
                    "Exemplo 1:\n"
                    "Colisão lateral entre dois veículos, em direções opostas, sendo o periciado atingido no flanco esquerdo.\n"
                    "Conclusão: Danos da esquerda para a direita (lataria entra no veículo) e de frente para trás (outro veículo vinha em direção contrária).\n"
                    "Exemplo 2:\n"
                    "Colisão frontal entre dois veículos, em direções opostas, sendo o periciado atingido na porção direita da frente.\n"
                    "Conclusão: Danos de frente para trás (lataria entra no veículo) e da direita para a esquerda (lataria sai do veículo).\n"
                    "Explique claramente quais vetores foram identificados no caso analisado, conforme os dados apresentados. Não deixe de incluir essas duas orientações de danos na conclusão.\n"
                ),
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                ],
            },
        ],
        max_tokens=350,
    )
    conclusao = response.choices[0].message.content.strip()

    conclusao_final = f"""## CONCLUSÃO

{conclusao}
"""
    return conclusao_final


def gerar_relatorio_completo(relatorio_nome):
    """Gera o relatório completo seguindo a nova estrutura"""
    # Carregar dados
    resultado = carregar_dados_relatorio(relatorio_nome)
    if resultado[0] is None:
        st.error(resultado[1])
        return None

    dados, descricoes, erro = resultado
    if erro:
        st.error(erro)
        return None

    # Verificar se descricoes é um dicionário válido
    if not isinstance(descricoes, dict):
        descricoes = {}

    # Carregar orientações da estrutura
    orientacoes = carregar_orientacoes_estrutura()

    # Gerar cada seção
    preambulo = gerar_preambulo(dados)
    historico_final = gerar_historico(dados, orientacoes)
    objetivo = gerar_objetivo(dados, orientacoes)
    quesitos = gerar_quesitos(dados)
    exames = gerar_exames(descricoes, orientacoes)
    conclusao_final = gerar_conclusao(dados, orientacoes)

    # Montar relatório completo
    relatorio_completo = f"""
# LAUDO PERICIAL - BO {dados["Número do BO"]}

{preambulo}

{historico_final}

{objetivo}

{quesitos}

{conclusao_final}

---
**Perito Responsável:** {dados["Nome do Perito"]}
**Data da Perícia:** {dados["Data da Perícia"]}
**BO:** {dados["Número do BO"]}

{exames}
"""

    return relatorio_completo


def salvar_relatorio(relatorio_nome, relatorio_completo):
    """Salva o relatório em formato TXT e DOCX"""
    pasta_relatorio = os.path.join(PASTA_LAUDOS, relatorio_nome)

    # Salvar TXT
    caminho_txt = os.path.join(pasta_relatorio, "relatorio_gerado.txt")
    with open(caminho_txt, "w", encoding="utf-8") as f:
        f.write(relatorio_completo)

    # Salvar DOCX
    try:
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading(
            f'LAUDO PERICIAL - BO {relatorio_nome.split("_")[1]}', 0)

        # Dividir o relatório em seções e adicionar ao documento
        secoes = relatorio_completo.split('\n\n')
        for secao in secoes:
            if secao.strip():
                if secao.startswith('#'):
                    # É um cabeçalho
                    nivel = secao.count('#')
                    texto = secao.replace('#', '').strip()
                    doc.add_heading(texto, min(nivel, 9))
                else:
                    doc.add_paragraph(secao)

        caminho_docx = os.path.join(pasta_relatorio, "relatorio_gerado.docx")
        doc.save(caminho_docx)

        return True, None
    except ImportError:
        return False, "Pacote python-docx não instalado. Apenas o arquivo TXT foi gerado."
    except Exception as e:
        return False, f"Erro ao gerar DOCX: {str(e)}"


# Interface principal
st.set_page_config(
    page_title="Geração de Relatório Pericial",
    page_icon="📋",
    layout="wide"
)

st.markdown("# 📋 Geração de Relatório Pericial")
st.markdown("""
<div style="background-color:#f0f4f8; border-radius:10px; padding: 1.5rem 1.5rem 1.2rem 1.5rem; border: 1px solid #dbeafe;">
    <h2 style="color:#2563eb; margin-bottom:0.5rem;">📝 Sistema de Geração de Laudos de Danos Veiculares</h2>
    <ul style="font-size:1.1rem; color:#222;">
        <li>O laudo gerado incluirá: <b>preâmbulo</b>, <b>histórico</b>, <b>quesitos</b> e uma <b>breve conclusão</b>.</li>
        <li>O <b>detalhamento dos danos</b> descritos pela IA, com base nas fotos carregadas, será apresentado ao final do documento.</li>
        <li><b>Importante:</b> Cabe ao perito avaliar a necessidade de uma descrição minuciosa dos danos, complementando o laudo conforme o caso.</li>
    </ul>
</div>
""", unsafe_allow_html=True)


# Verificar se as orientações da estrutura estão disponíveis
orientacoes_disponiveis = carregar_orientacoes_estrutura()
if not orientacoes_disponiveis:
    st.warning(
        "⚠️ **Atenção:** As orientações da estrutura do laudo não foram encontradas.")
    st.info("Acesse a página '06_EstruturaLaudo' para configurar as orientações antes de gerar relatórios.")
    st.stop()

# st.success("✅ Orientações da estrutura do laudo carregadas com sucesso!")

# Sidebar para seleção de relatório
st.sidebar.markdown("## 📁 Seleção de Relatório")

relatorios_existentes = listar_relatorios()
if not relatorios_existentes:
    st.sidebar.warning("Nenhum relatório encontrado.")
    st.info("Crie um relatório na página '01_CriaRelatorio' primeiro.")
    st.stop()

relatorio_selecionado = st.sidebar.selectbox(
    "Selecionar Relatório",
    ["Nenhum"] + relatorios_existentes,
    help="Selecione o relatório para o qual deseja gerar o laudo"
)

if relatorio_selecionado != "Nenhum":
    st.sidebar.markdown(f"### 📋 Relatório {relatorio_selecionado}")

    # Verificar status do relatório
    relatorio_existe, info_relatorio = verificar_relatorio_existente(
        relatorio_selecionado)

    if relatorio_existe:
        st.sidebar.success("✅ Laudo já foi gerado!")

        # Informações do relatório existente
        with st.sidebar.expander("📊 Informações do Laudo"):
            if 'txt' in info_relatorio:
                st.write(f"**Arquivo TXT:** ✅")
                st.write(f"Tamanho: {info_relatorio['txt']['tamanho']} bytes")
                st.write(
                    f"Última modificação: {info_relatorio['txt']['data_modificacao'].strftime('%d/%m/%Y %H:%M:%S')}")

            if 'docx' in info_relatorio:
                st.write(f"**Arquivo DOCX:** ✅")
                st.write(f"Tamanho: {info_relatorio['docx']['tamanho']} bytes")
                st.write(
                    f"Última modificação: {info_relatorio['docx']['data_modificacao'].strftime('%d/%m/%Y %H:%M:%S')}")

        # Opções para relatório existente
        col1, col2 = st.sidebar.columns(2)
        with col1:
            if st.button("📖 Visualizar Laudo", type="secondary"):
                st.session_state['visualizar_laudo'] = True
        with col2:
            if st.button("🔄 Nova Versão", type="primary"):
                st.session_state['gerar_nova_versao'] = True
                st.session_state['visualizar_laudo'] = False

        # Download dos arquivos
        st.sidebar.markdown("### 📥 Download dos Arquivos")
        pasta_relatorio = os.path.join(PASTA_LAUDOS, relatorio_selecionado)

        if os.path.exists(os.path.join(pasta_relatorio, "relatorio_gerado.txt")):
            with open(os.path.join(pasta_relatorio, "relatorio_gerado.txt"), "r", encoding="utf-8") as f:
                txt_content = f.read()
            st.sidebar.download_button(
                label="📄 Baixar TXT",
                data=txt_content,
                file_name=f"laudo_{relatorio_selecionado}.txt",
                mime="text/plain"
            )

        if os.path.exists(os.path.join(pasta_relatorio, "relatorio_gerado.docx")):
            with open(os.path.join(pasta_relatorio, "relatorio_gerado.docx"), "rb") as f:
                docx_content = f.read()
            st.sidebar.download_button(
                label="📝 Baixar DOCX",
                data=docx_content,
                file_name=f"laudo_{relatorio_selecionado}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    else:
        st.sidebar.info("ℹ️ Laudo ainda não foi gerado.")
        if st.sidebar.button("🚀 Gerar Laudo", type="primary"):
            st.session_state['gerar_laudo'] = True

# Controle de estado para mostrar conteúdo
if relatorio_selecionado != "Nenhum":
    # Verificar se deve gerar laudo
    if st.session_state.get('gerar_laudo', False) or st.session_state.get('gerar_nova_versao', False):
        st.markdown("## 🚀 Gerando Laudo Pericial...")

        with st.spinner("Gerando laudo com base nas orientações da estrutura..."):
            relatorio_completo = gerar_relatorio_completo(
                relatorio_selecionado)

            if relatorio_completo:
                # Salvar arquivos
                sucesso, erro = salvar_relatorio(
                    relatorio_selecionado, relatorio_completo)

                if sucesso:
                    st.success("✅ Laudo pericial gerado com sucesso!")
                    st.info(
                        "Arquivos TXT e DOCX foram criados na pasta do relatório.")
                else:
                    st.warning(f"⚠️ {erro}")

                # Exibir preview do laudo
                st.markdown("## 📋 Preview do Laudo Gerado")
                st.text_area("Laudo Completo", relatorio_completo,
                             height=600, disabled=True)

                # Botão para visualizar
                if st.button("👁️ Visualizar Laudo Completo"):
                    st.session_state['visualizar_laudo'] = True
                    st.rerun()
            else:
                st.error("❌ Erro ao gerar o laudo pericial.")

        # Reset dos estados
        st.session_state['gerar_laudo'] = False
        st.session_state['gerar_nova_versao'] = False

    # Verificar se deve visualizar laudo existente
    elif st.session_state.get('visualizar_laudo', False):
        st.markdown("## 📖 Visualizando Laudo Existente")

        pasta_relatorio = os.path.join(PASTA_LAUDOS, relatorio_selecionado)
        caminho_txt = os.path.join(pasta_relatorio, "relatorio_gerado.txt")

        if os.path.exists(caminho_txt):
            with open(caminho_txt, "r", encoding="utf-8") as f:
                conteudo_laudo = f.read()

            st.text_area("Laudo Completo", conteudo_laudo,
                         height=600, disabled=True)
        else:
            st.error("Arquivo do laudo não encontrado.")

        # Botão para voltar
        if st.button("🔙 Voltar"):
            st.session_state['visualizar_laudo'] = False
            st.rerun()

    # Estado inicial - mostrar informações básicas
    else:
        st.markdown("## 📋 Geração de Laudo Pericial")
        st.info("Selecione uma das opções na barra lateral para prosseguir.")

        # Mostrar informações do relatório selecionado
        resultado = carregar_dados_relatorio(relatorio_selecionado)
        if resultado[0] is not None:
            dados, descricoes, _ = resultado

            # Verificar se descricoes é um dicionário válido
            if not isinstance(descricoes, dict):
                descricoes = {}

            # Contar elementos
            num_imagens = len(
                [k for k in descricoes.keys() if k != "ditado_pericia"])
            num_ditados = len(descricoes.get("ditado_pericia", []))

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("📸 Imagens", num_imagens)
            with col2:
                st.metric("🎤 Ditados", num_ditados)
            with col3:
                st.metric(
                    "📁 Status", "✅ Pronto" if relatorio_existe else "⏳ Pendente")
            with col4:
                st.metric(
                    "📋 Estrutura", "✅ Configurada" if orientacoes_disponiveis else "❌ Pendente")

            # Mostrar orientações ativas
            st.markdown("### 🎯 Orientações da Estrutura Ativas")
            with st.expander("Ver orientações configuradas"):
                for chave, valor in orientacoes_disponiveis.items():
                    if chave not in ['data_criacao', 'versao']:
                        st.markdown(f"**{chave.replace('_', ' ').title()}:**")
                        st.text(valor[:200] + "..." if len(valor)
                                > 200 else valor)
        else:
            st.error(
                f"Erro ao carregar informações do relatório: {resultado[1]}")

else:
    st.info("👈 Selecione um relatório na barra lateral para começar.")
